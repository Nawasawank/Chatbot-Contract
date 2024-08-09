from database import db, Rent, State  
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from uploadfile import upload_to_fileio
from bahttext import bahttext
from pythainlp.tag import NER
import re
from linebot.v3.messaging import  QuickReply, QuickReplyItem, PostbackAction

ner = NER("thainer")

def CheckPlace(text):
    output_ner = ner.tag(text)
    place = ""
    for element in output_ner:
        if element[1] in ["B-LOCATION", "I-LOCATION"]:
            place += element[0]
    return place

def CheckPerson(text):
    ner = NER()
    output_ner = ner.tag(text)
    name = ""
    for element in output_ner:
        if element[1] in ["B-PERSON", "I-PERSON"]:
            name += element[0]
    return name

def CheckDate(text):
    ner = NER()
    output_ner = ner.tag(text)
    date = ""
    for element in output_ner:
        if element[1] in ["B-DATE", "I-DATE"]:
            date += element[0]
    return date

def Check_NumberAddress(text):
    pattern = r'^\d{1,4}(\/\d{1,4})?$'
    return bool(re.match(pattern, text))

def Check_IDcard(text):
    pattern = r'^\d{13}$'
    return bool(re.match(pattern, text))

def rental_contract(sender_id, text):
    rent_entry = db.session.query(Rent).filter_by(sender_id=sender_id).order_by(Rent.id.desc()).first()
    curstate = State.query.filter_by(sender_id=sender_id).first()
    current_state_value = curstate.current_state
    print("Current state value:", current_state_value)
    
    entities = ner.tag(text)
    print("Named Entities:", entities)
    
    response = ""
    quick_reply = None
    
    if current_state_value == "1":
        curstate.type_contract = "rental"
        curstate.current_state = "2"
        db.session.commit()
        response = "กรุณากรอกสถานที่ทำสัญญา"
    elif current_state_value == "2":
        place = CheckPlace(text)
        if not place:
            response = "กรุณาระบุสถานที่ให้ถูกต้อง"
        else:
            rent_entry = Rent(sender_id=sender_id, place=place)
            db.session.add(rent_entry)
            db.session.commit()
            curstate.current_state = "3"
            db.session.commit()
            response = "กรุณากรอกชื่อผู้ให้เช่า"
    elif current_state_value == "3":
        name1 = CheckPerson(text)
        if not name1:
            response = "กรุณากรอกชื่อผู้ให้เช่าที่ถูกต้อง"
        else:
            rent_entry.name1 = name1
            db.session.commit()
            curstate.current_state = "4"
            db.session.commit()
            response = "กรุณากรอกที่อยู่อำเภอผู้ให้เช่า"
    elif current_state_value == "4":
        place = CheckPlace(text)
        if not place:
            response = "กรุณาระบุอำเภอที่ให้ถูกต้อง"
        else:
            rent_entry.district1 = text
            db.session.commit()
            curstate.current_state = "5"
            db.session.commit()
            response = "กรุณากรอกจังหวัด"
    elif current_state_value == "5":
        province = CheckPlace(text)
        if not province:
            response = "กรุณาระบุจังหวัดให้ถูกต้อง"
        else:
            rent_entry.province1 = text
            db.session.commit()
            curstate.current_state = "6"
            db.session.commit()
            response = "กรุณากรอกชื่อผู้เช่า"
    elif current_state_value == "6":
        name2 = CheckPerson(text)
        if not name2:
            response = "กรุณากรอกชื่อผู้เช่าให้ถูกต้อง"
        else:
            rent_entry.name2 = text
            db.session.commit()
            curstate.current_state = "7"
            db.session.commit()
            response = "กรุณากรอกเลขบัตรประชาชน"
    elif current_state_value == "7":
        idcard2 = Check_IDcard(text)
        if not idcard2:
            response = "กรุณากรอกเลขบัตรประชาชนให้ถูกต้อง"
        else:
            rent_entry.idcard2 = text
            db.session.commit()
            curstate.current_state = "8"
            db.session.commit()
            response = "กรุณากรอกอายุผู้เช่า"
    elif current_state_value == "8":
        if text.isnumeric(): 
            rent_entry.age2 = text
            db.session.commit()
            curstate.current_state = "9"
            db.session.commit()
            response = "กรุณากรอกบ้านเลขที่"
        else:
            response = "กรุณากรอกอายุเป็นตัวเลข"
    elif current_state_value == "9":
        no_address = Check_NumberAddress(text)
        if not no_address:
            response = "กรุณากรอกบ้านเลขที่ให้ถูกต้อง"
        else:
            rent_entry.house2 = text
            db.session.commit()
            curstate.current_state = "10"
            db.session.commit()
            response = "กรุณากรอกหมู่"
    elif current_state_value == "10":
        if text.isnumeric(): 
            rent_entry.vilno2 = text
            db.session.commit()
            curstate.current_state = "11"
            db.session.commit()
            response = "ถนน"
        else:
            response = "กรุณากรอกหมู่ให้ถูกต้อง"
    elif current_state_value == "11":
        road = CheckPlace(text)
        if not road:
            response = "กรุณากรอกถนนให้ถูกต้อง"
        else:
            rent_entry.street2 = text
            db.session.commit()
            curstate.current_state = "12"
            db.session.commit()
            response = "ซอย"
    elif current_state_value == "12":
        soi = CheckPlace(text)
        if not soi:
            response = "กรุณากรอกซอยให้ถูกต้อง"
        else:
            rent_entry.lane2 = text
            db.session.commit()
            curstate.current_state = "13"
            db.session.commit()
            response = "ตำบล/แขวง"
    elif current_state_value == "13":
        Subdistrict = CheckPlace(text)
        if not Subdistrict:
            response = "กรุณากรอกตำบลให้ถูกต้อง"
        else:
            rent_entry.subd2 = text
            db.session.commit()
            curstate.current_state = "14"
            db.session.commit()
            response = "อำเภอ/เขต"
    elif current_state_value == "14":
        district = CheckPlace(text)
        if not district:
            response = "กรุณากรอกอำเภอให้ถูกต้อง"
        else:
            rent_entry.district2 = text
            db.session.commit()
            curstate.current_state = "15"
            db.session.commit()
            response = "จังหวัด"
    elif current_state_value == "15":
        province = CheckPlace(text)
        if not province:
            response = "กรุณากรอกจังหวัดให้ถูกต้อง"
        else:
            rent_entry.province2 = text
            db.session.commit()
            curstate.current_state = "16"
            db.session.commit()
            response = "กรุณากรอกชื่อผู้ออกบัตรประชาชน"
    elif current_state_value == "16":
        name_card = CheckPerson(text)
        if not name_card:
            response = "กรุณากรอกชื่อผู้ออกบัตรให้ถูกต้อง"
        else:
            rent_entry.authority = text
            db.session.commit()
            curstate.current_state = "17"
            db.session.commit()
            response = "วันที่ออกบัตร"
    elif current_state_value == "17":
        date = CheckDate(text)
        if not date:
            response = "กรุณากรอกวันที่ให้ถูกต้อง"
        else:
            rent_entry.dateofid = text
            db.session.commit()
            curstate.current_state = "18"
            db.session.commit()

            quick_reply = QuickReply(items=[
                QuickReplyItem(action=PostbackAction(label="บ้าน", data="property_house")),
                QuickReplyItem(action=PostbackAction(label="ที่ดิน", data="property_land"))
            ])
            response = "กรุณากรอกสิ่งที่เช่า"
    elif current_state_value == "18":
        if text == "property_house":
            rent_entry.property = "บ้าน"
            db.session.commit()
            curstate.current_state = "19"
            db.session.commit()
            response = "กรุณากรอกรายละเอียดบ้าน"
        elif text == "property_land":
            rent_entry.property = "ที่ดิน"
            db.session.commit()
            curstate.current_state = "19"
            db.session.commit()
            response = "กรุณากรอกรายละเอียดที่ดิน"
    elif current_state_value == "19":
        rent_entry.detail_property = text
        db.session.commit()
        curstate.current_state = "20"
        db.session.commit()
        response = "กรุณากรอกวัตถุประสงค์ในการเช่า"
    elif current_state_value == "20":
        rent_entry.purpose = text
        db.session.commit()
        curstate.current_state = "21"
        db.session.commit()
        response = "ระยะเวลาในการเช่า (ex. 3 ปี)"
    elif current_state_value == "21":
        rent_entry.duration = text
        db.session.commit()
        curstate.current_state = "22"
        db.session.commit()
        response = "วันที่เริ่มเช่า"
    elif current_state_value == "22":
        fromdate = CheckDate(text)
        if not fromdate:
            response = "กรุณากรอกวันที่เริ่มเช่าให้ถูกต้อง"
        else:
            rent_entry.fromdate = text
            db.session.commit()
            curstate.current_state = "23"
            db.session.commit()
            response = "วันที่สิ้นสุดการเช่า"
    elif current_state_value == "23":
        todate = CheckDate(text)
        if not todate:
            response = "กรุณากรอกวันที่สิ้นสุดการเช่าให้ถูกต้อง"
        else:
            rent_entry.todate = text
            db.session.commit()
            curstate.current_state = "24"
            db.session.commit()
            response = "ประเภทการจ่ายค่าเช่า (รายวัน/รายเดือน/รายปี)"
    elif current_state_value == "24":
        rent_entry.typeofrent = text
        db.session.commit()
        curstate.current_state = "25"
        db.session.commit()
        response = "จำนวนเงินค่าเช่า"
    elif current_state_value == "25":
        rent_entry.price = text
        db.session.commit()
        curstate.current_state = "26"
        db.session.commit()
        response = "วันที่กำหนดชำระ"
    elif current_state_value == "26":
        rent_entry.duedate = text
        db.session.commit()
        curstate.current_state = "27"
        db.session.commit()
        response = "กรุณากรอกผู้เสียภาษีจากการเช่า"
    elif current_state_value == "27":
        rent_entry.tax = text
        db.session.commit()
        file_link = generate_document_rent(sender_id)
        db.session.delete(curstate)
        db.session.commit()
        response = f"คลิกที่ลิงก์เพื่อดาวน์โหลดไฟล์: {file_link}"

    if quick_reply:
        return response, quick_reply
    else:
        return response, None

def generate_document_rent(sender_id):
    rent_entry = db.session.query(Rent).filter_by(sender_id=sender_id).order_by(Rent.id.desc()).first()
    if not rent_entry:
        return None

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    style = doc.styles['Normal']
    style.font.name = 'Angsana New'
    style.font.size = Pt(14)

    heading = doc.add_heading('', level=1)
    run = heading.add_run('หนังสือสัญญาเช่า')
    run.font.name = 'Angsana New'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    place_date = details.add_run(f"เขียนที่ {rent_entry.place}\n    เมื่อวันที่ {rent_entry.date.strftime('%d/%m/%Y')}")
    place_date.bold = False
    place_date.font.size = Pt(14)
    details.paragraph_format.line_spacing = Pt(16)

    contract_parties = doc.add_paragraph()
    contract_parties.add_run(f"โดยหนังสือฉบับนี้ ข้าพเจ้า {rent_entry.name1} ")
    contract_parties.add_run(f"อำเภอ {rent_entry.district1} จังหวัด {rent_entry.province1} ")
    contract_parties.add_run('ซึ่งต่อไปในสัญญานี้เรียกว่า "ผู้ให้เช่า" ฝ่ายหนึ่ง ')
    contract_parties.add_run(f"กับข้าพเจ้า {rent_entry.name2} เลขประจำตัวประชาชน { rent_entry.idcard2} อายุ {rent_entry.age2} ปี ")
    contract_parties.add_run(f"อยู่บ้านเลขที่ {rent_entry.house2} หมู่ที่ {rent_entry.vilno2} ")
    contract_parties.add_run(f"ถนน {rent_entry.street2} ตรอก/ซอย {rent_entry.lane2} ")
    contract_parties.add_run(f"ตำบล {rent_entry.subd2} อำเภอ {rent_entry.district2} ")
    contract_parties.add_run(f"จังหวัด {rent_entry.province2} ")
    contract_parties.add_run('ซึ่งต่อไปในสัญญานี้เรียกว่า "ผู้เช่า" อีกฝ่ายหนึ่ง ทั้งสองฝ่ายตกลงทำสัญญากันมีข้อความต่อไปนี้ คือ')
    contract_parties.paragraph_format.line_spacing = Pt(18)
    contract_parties.paragraph_format.space_before = Pt(0)
    contract_parties.paragraph_format.space_after = Pt(0)
    contract_parties.paragraph_format.first_line_indent = Pt(36)

    thai_text_price = bahttext(int(rent_entry.price))

    first = doc.add_paragraph()
    first_head = first.add_run(f"ข้อ 1. ")
    first_head.bold = True
    first.add_run(f"ผู้ให้เช่าตกลงให้เช่าและผู้เช่าตกลงเช่า {rent_entry.property} โดยมีรายละเอียดดังนี้ {rent_entry.detail_property} ")
    first.add_run(f"มีวัตถุประสงค์เพื่อ {rent_entry.purpose} ")
    first.add_run(f"มีกำหนดเวลาเช่า {rent_entry.duration} (วัน/เดือน/ปี) ตั้งแต่วันที่ {rent_entry.fromdate} ถึงวันที่ {rent_entry.todate} ")
    first.add_run(f"โดยผู้เช่าตกลงจ่ายค่าเช่าเป็นราย (วัน/เดือน/ปี) {rent_entry.typeofrent} ละ {rent_entry.price} บาท ({thai_text_price}) ")
    first.add_run(f"มีกำหนดชำระวันที่ {rent_entry.duedate} ของทุก ๆ เดือน ส่วนเงินค่าภาษี อันเกิดจากทรัพย์สินที่เช่านี้ให้ {rent_entry.tax} เป็นผู้เสีย")
    first.paragraph_format.line_spacing = Pt(18)
    first.paragraph_format.first_line_indent = Pt(36)
    first.paragraph_format.space_before = Pt(0)
    first.paragraph_format.space_after = Pt(0)

    second = doc.add_paragraph()
    second_head = second.add_run(f"ข้อ 2. ")
    second_head.bold = True
    second.add_run('ผู้เช่าได้ตรวจดูทรัพย์สินที่เช่าแล้ว เห็นว่าทุกสิ่งอยูในสภาพเรียบร้อยใช้การได้อย่างสมบูรณ์จะดูแลทรัพย์สิน ที่เช่า มิได้ให้สูญหายและบำรุงรักษาให้อยูในสภาพดีอยู่เสมอ พร้อมที่จะส่งมอบคืน ตามสภาพเดิมทุกประการและตกลงยอมให้ ผู้ให้เช่าหรือตัวแทน เข้าตรวจดูทรัพย์สินที่เช่าได้ทุกเวลา ภายหลังที่ได้แจ้งความประสงค์ให้ผู้เช่าทราบแล้ว')
    second.paragraph_format.line_spacing = Pt(18)
    second.paragraph_format.first_line_indent = Pt(36)
    second.paragraph_format.space_before = Pt(0)
    second.paragraph_format.space_after = Pt(0)

    third = doc.add_paragraph()
    third_head = third.add_run(f"ข้อ 3. ")
    third_head.bold = True
    third.add_run('ผู้เช่าไม่มีสิทธินำทรัพย์สินที่เช่าออกให้ผู้อื่นเช่าช่วงหรือทำนิติกรรมใด ๆ กับผู้อื่นในอันที่จะเป็นผล ก่อให้เกิดความผูกพันในทรัพย์สินที่เช่า ไม่ว่าโดยตรงหรือโดยปริยายและจะไม่ทำการดัดแปลง หรือต่อเติมทรัพย์สิน ที่เช่าไม่ว่าทั้งหมดหรือบางส่วน เว้นแต่จะได้รับความยินยอม เป็นหนังสือจากผู้ให้เช่าแกละหากผู้เช่าได้ทำการดัดแปลงหรือต่อเติมสิ่งใดตามที่ได้รับความยินยอมเมื่อใดแล้ว ผู้เช่ายอมยกกรรมสิทธิ์ในทรัพย์สินสิ่งนั้น ให้ตกเป็นของผู้ให้เช่านับแต่เมื่อนั้นด้วยทั้งสิ้น')
    third.paragraph_format.line_spacing = Pt(18)
    third.paragraph_format.first_line_indent = Pt(36)
    third.paragraph_format.space_before = Pt(0)
    third.paragraph_format.space_after = Pt(0)

    fourth = doc.add_paragraph()
    fourth_head = fourth.add_run(f"ข้อ 4. ")
    fourth_head.bold = True
    fourth.add_run('เมื่อผู้เช่ากระทำผิดสัญญาข้อหนึ่งข้อใด ผู้ให้เช่ามีสิทธิบอกเลิกสัญญาได้ทันที และผู้เช่ายอมชดใช้ ค่าฤชาธรรมเนียม กับค่าทนายความจนค่าพาหนะและค่าใช้จ่ายในการติดตามทวงถามให้แก่ผู้ให้เช่าจนครบถ้วนหากมีความเสียหายดังกล่าวเกิดขึ้นเพราะผู้เช่าเป็นฝ่ายผิดสัญญา')
    fourth.paragraph_format.line_spacing = Pt(18)
    fourth.paragraph_format.first_line_indent = Pt(36)
    fourth.paragraph_format.space_before = Pt(0)
    fourth.paragraph_format.space_after = Pt(0)

    last = doc.add_paragraph()
    last.add_run('คู่สัญญาได้อ่านและเข้าใจข้อความดีแล้ว จึงลงลายมือชื่อไว้เป็นสำคัญต่อหน้าพยาน')
    last.paragraph_format.line_spacing = Pt(18)
    last.paragraph_format.first_line_indent = Pt(36)
    last.paragraph_format.space_before = Pt(0)
    last.paragraph_format.space_after = Pt(0)

    spacing_paragraph = doc.add_paragraph()
    spacing_paragraph.paragraph_format.space_before = Pt(0)
    spacing_paragraph.paragraph_format.space_after = Pt(10)

    signature_table = doc.add_table(rows=3, cols=2)
    signature_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in signature_table.rows:
        for cell in row.cells:
            cell.width = Pt(200)

    cell_0_0 = signature_table.cell(0, 0).paragraphs[0]
    cell_0_0.add_run('ลงชื่อ..............................................ผู้ให้เช่า\n')
    cell_0_0.add_run('        (..............................................)').font.size = Pt(14)

    cell_0_1 = signature_table.cell(0, 1).paragraphs[0]
    cell_0_1.add_run('ลงชื่อ..............................................ผู้เช่า\n')
    cell_0_1.add_run('        (..............................................)').font.size = Pt(14)

    cell_1_0 = signature_table.cell(1, 0).paragraphs[0]
    cell_1_0.add_run('ลงชื่อ..............................................พยาน\n')
    cell_1_0.add_run('        (..............................................)').font.size = Pt(14)

    cell_1_1 = signature_table.cell(1, 1).paragraphs[0]
    cell_1_1.add_run('ลงชื่อ..............................................พยาน\n')
    cell_1_1.add_run('        (..............................................)').font.size = Pt(14)

    file_path = 'contract.docx'
    doc.save(file_path)
    return upload_to_fileio(file_path)
