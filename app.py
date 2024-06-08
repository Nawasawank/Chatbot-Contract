from flask import Flask, request, send_file
from pymessenger import Bot
from docx import Document
import requests
from dotenv import load_dotenv
from os import environ
from flask_sqlalchemy import SQLAlchemy
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from datetime import datetime

load_dotenv()

app = Flask(__name__)

VERIFY_TOKEN = environ.get("VERIFY_TOKEN")
PAGE_ACCESS_TOKEN = environ.get("PAGE_ACCESS_TOKEN")
print("Verify Token:", VERIFY_TOKEN)
print("Page Access Token:", PAGE_ACCESS_TOKEN)
bot = Bot(PAGE_ACCESS_TOKEN)


# Construct the database URI
POSTGRES_HOST = environ.get("POSTGRES_HOST")
POSTGRES_USER = environ.get("POSTGRES_USER")
POSTGRES_PASSWORD = environ.get("POSTGRES_PASSWORD")
POSTGRES_DATABASE = environ.get("POSTGRES_DATABASE")

SQLALCHEMY_DATABASE_URI = f'postgresql://{POSTGRES_USER}:{POSTGRES_PASSWORD}@{POSTGRES_HOST}/{POSTGRES_DATABASE}'

app.config['SQLALCHEMY_DATABASE_URI'] = SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)


class Rent(db.Model):
    __tablename__ = 'rent'
    id = db.Column(db.Integer, primary_key=True)
    sender_id = db.Column(db.String(255), nullable=False)
    place = db.Column(db.String(255), nullable=False)
    date = db.Column(db.Date, default=db.func.current_date())
    name1 = db.Column(db.String(255), nullable=False)
    district1 = db.Column(db.String(255))
    province1 = db.Column(db.String(255))
    name2 = db.Column(db.String(255), nullable=False)
    age2 = db.Column(db.String(255))
    house2 = db.Column(db.String(255))
    vilno2 = db.Column(db.String(255))
    street2 = db.Column(db.String(255))
    lane2 = db.Column(db.String(255))
    subd2 = db.Column(db.String(255))
    district2 = db.Column(db.String(255))
    province2 = db.Column(db.String(255))
    idcard2 = db.Column(db.String(255))
    authority = db.Column(db.String(255))
    dateofid = db.Column(db.String(255))
    property = db.Column(db.String(255))
    purpose = db.Column(db.String(255))
    fromdate = db.Column(db.String(255))
    todate = db.Column(db.String(255))
    typeofrent = db.Column(db.String(255))
    price = db.Column(db.String(255))
    duedate = db.Column(db.String(255))
    tax = db.Column(db.String(255))


class State(db.Model):
    __tablename__ = 'current_state'
    id = db.Column(db.Integer, primary_key=True)
    sender_id = db.Column(db.String(100), nullable=False)
    current_state = db.Column(db.String(100), nullable=False)
    type_contract = db.Column(db.String(100), nullable=False)


def verify_token(req):
    token = req.args.get("hub.verify_token")
    print("Received Token:", token)
    if token == VERIFY_TOKEN:
        print("Token was verified")
        return req.args.get("hub.challenge")
    else:
        print("Token was not verified")
        return "incorrect"


@app.route("/webhook", methods=["GET", "POST"])
def listen():
    if request.method == "GET":
        return verify_token(request)
    if request.method == 'POST':
        payload = request.json
        print("Payload received:", payload)
        for entry in payload.get("entry", []):
            for messaging_event in entry.get("messaging", []):
                if "message" in messaging_event and "text" in messaging_event["message"]:
                    sender_id = messaging_event["sender"]["id"]
                    message_text = messaging_event["message"]["text"]
                    response = process_message(message_text, sender_id)
                    if response:
                        bot.send_text_message(sender_id, response)
                else:
                    print("Invalid message format. Skipping.")
        return "ok"


def process_message(message_text, sender_id):
    text = message_text
    curstate = State.query.filter_by(sender_id=sender_id).first()
    if not curstate:  # If sender_id is not found in current_state table
        # Create a new entry for sender_id
        curstate = State(sender_id=sender_id, current_state="1", type_contract="")
        db.session.add(curstate)
        db.session.commit()
    else:
        if curstate.type_contract == "rental":
            return rental_contract(sender_id, text)
        # elif type == "sale":
        #    return sale_contract(sender_id, text, conversation_state)
    if text == "สัญญาเช่า":
        return rental_contract(sender_id, text)
    # elif text == "สัญญาขาย":
    #    return sale_contract(sender_id, text, conversation_state)
    else:
        return "ประเภทของสัญญาไม่ถูกต้อง"


def rental_contract(sender_id, text):
    rent_entry = db.session.query(Rent).filter_by(sender_id=sender_id).order_by(Rent.id.desc()).first()
    curstate = State.query.filter_by(sender_id=sender_id).first()
    current_state_value = curstate.current_state
    print("Current state value:", current_state_value)
    if current_state_value == "1":
        curstate.type_contract = "rental"
        curstate.current_state = "2"
        db.session.commit()
        return "กรุณากรอกสถานที่ทำสัญญา"
    elif current_state_value == "2":
        rent_entry = Rent(sender_id=sender_id, place=text)
        db.session.add(rent_entry)
        db.session.commit()
        curstate.current_state = "3"
        db.session.commit()
        return "กรุณากรอกชื่อผู้ให้เช่า"
    elif current_state_value == "3":
        if rent_entry:
            rent_entry.name1 = text
            db.session.commit()
            curstate.current_state = "4"
            db.session.commit()
            return "กรุณากรอกที่อยู่อำเภอผู้ให้เช่า"
        else:
            return "Rent entry not found for the provided sender_id."
    elif current_state_value == "4":
        if rent_entry:
            rent_entry.district1 = text
            db.session.commit()
            curstate.current_state = "5"
            db.session.commit()
            return "กรุณากรอกจังหวัด"
    elif current_state_value == "5":
        if rent_entry:
            rent_entry.province1 = text
            db.session.commit()
            curstate.current_state = "6"
            db.session.commit()
            return "กรุณากรอกชื่อผู้เช่า"
    elif current_state_value == "6":
        if rent_entry:
            rent_entry.name2 = text
            db.session.commit()
            curstate.current_state = "7"
            db.session.commit()
            return "กรุณากรอกอายุผู้เช่า"
    elif current_state_value == "7":
        if rent_entry:
            rent_entry.age2 = text
            db.session.commit()
            curstate.current_state = "8"
            db.session.commit()
            return "กรุณากรอกบ้านเลขที่"
    elif current_state_value == "8":
        if rent_entry:
            rent_entry.house2 = text
            db.session.commit()
            curstate.current_state = "9"
            db.session.commit()
            return "กรุณากรอกหมู่"
    elif current_state_value == "9":
        if rent_entry:
            rent_entry.vilno2 = text
            db.session.commit()
            curstate.current_state = "10"
            db.session.commit()
            return "ถนน"
    elif current_state_value == "10":
        if rent_entry:
            rent_entry.street2 = text
            db.session.commit()
            curstate.current_state = "11"
            db.session.commit()
            return "ซอย"
    elif current_state_value == "11":
        if rent_entry:
            rent_entry.lane2 = text
            db.session.commit()
            curstate.current_state = "12"
            db.session.commit()
            return "ตำบล/แขวง"
    elif current_state_value == "12":
        if rent_entry:
            rent_entry.subd2 = text
            db.session.commit()
            curstate.current_state = "13"
            db.session.commit()
            return "อำเภอ/เขต"
    elif current_state_value == "13":
        if rent_entry:
            rent_entry.district2 = text
            db.session.commit()
            curstate.current_state = "14"
            db.session.commit()
            return "จังหวัด"
    elif current_state_value == "14":
        if rent_entry:
            rent_entry.province2 = text
            db.session.commit()
            curstate.current_state = "15"
            db.session.commit()
            return "เลขบัตรประชาชน"
    elif current_state_value == "15":
        if rent_entry:
            rent_entry.idcard2 = text
            db.session.commit()
            curstate.current_state = "16"
            db.session.commit()
            return "ชื่อผู้ออกบัตร"
    elif current_state_value == "16":
        if rent_entry:
            rent_entry.authority = text
            db.session.commit()
            curstate.current_state = "17"
            db.session.commit()
            return "วันที่ออกบัตร"
    elif current_state_value == "17":
        if rent_entry:
            rent_entry.dateofid = text
            db.session.commit()
            curstate.current_state = "18"
            db.session.commit()
            return "กรุณากรอกสิ่งที่เช่า"
    elif current_state_value == "18":
        if rent_entry:
            rent_entry.property = text
            db.session.commit()
            curstate.current_state = "19"
            db.session.commit()
            return "กรุณากรอกวัตถุประสงค์ในการเช่า"
    elif current_state_value == "19":
        if rent_entry:
            rent_entry.purpose = text
            db.session.commit()
            curstate.current_state = "20"
            db.session.commit()
            return "วันที่เริ่มเช่า"
    elif current_state_value == "20":
        if rent_entry:
            rent_entry.fromdate = text
            db.session.commit()
            curstate.current_state = "21"
            db.session.commit()
            return "วันที่สิ้นสุดการเช่า"
    elif current_state_value == "21":
        if rent_entry:
            rent_entry.todate = text
            db.session.commit()
            curstate.current_state = "22"
            db.session.commit()
            return "ประเภทการจ่ายค่าเช่า(รายวัน/รายเดือน/รายปี)"
    elif current_state_value == "22":
        if rent_entry:
            rent_entry.typeofrent = text
            db.session.commit()
            curstate.current_state = "23"
            db.session.commit()
            return "จำนวนเงินค่าเช่า"
    elif current_state_value == "23":
        if rent_entry:
            rent_entry.price = text
            db.session.commit()
            curstate.current_state = "24"
            db.session.commit()
            return "วันที่กำหนดชำระ"
    elif current_state_value == "24":
        if rent_entry:
            rent_entry.duedate = text
            db.session.commit()
            curstate.current_state = "25"
            db.session.commit()
            return "กรุณากรอกผู้เสียภาษีจากการเช่า"
    elif current_state_value == "25":
        if rent_entry:
            rent_entry.tax = text
            db.session.commit()
            file_link = generate_document_rent(sender_id)
            db.session.delete(curstate)
            db.session.commit()
            return f"คลิกที่ลิงก์เพื่อดาวน์โหลดไฟล์: {file_link}"


def generate_document_rent(sender_id):
    curstate = State.query.filter_by(sender_id=sender_id).first()
    if curstate and curstate.type_contract == "rental":
        rent_entry = Rent.query.filter_by(sender_id=sender_id).order_by(Rent.id.desc()).first()
        place = rent_entry.place
        date = rent_entry.date
        name1 = rent_entry.name1
        district1 = rent_entry.district1
        # Add other necessary fields

    doc = Document()
    doc.add_heading('หนังสือสัญญาเช่า', 0)
    doc.add_paragraph(f'เขียนที่ {place}')
    doc.add_paragraph(f'เมื่อวันที่ {date}')
    # Add other necessary fields

    file_path = 'contract.docx'
    doc.save(file_path)
    return upload_to_fileio(file_path)


def upload_to_fileio(file_path):
    with open(file_path, 'rb') as file:
        response = requests.post('https://file.io/', files={'file': file})
        response_json = response.json()
        file_link = response_json.get('link')
    return file_link


@app.route("/download/<path:filename>", methods=["GET"])
def download_file(filename):
    return send_file(filename, as_attachment=True)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=3000)










